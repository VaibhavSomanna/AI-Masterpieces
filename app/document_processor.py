import os
import uuid
from pathlib import Path
from typing import List, Dict, Any
import asyncio
from concurrent.futures import ThreadPoolExecutor

import PyPDF2
import pypdf
from docx import Document as DocxDocument
import pandas as pd
import tabula
import cv2
import pytesseract
from PIL import Image
from loguru import logger

from app.models import ExtractedData, DocumentChunk, ContentType

class DocumentProcessor:
    def __init__(self):
        self.executor = ThreadPoolExecutor(max_workers=4)
        
    async def process_document(self, file_path: str) -> ExtractedData:
        """Process a document and extract text, tables, and images"""
        file_path = Path(file_path)
        file_extension = file_path.suffix.lower()
        
        logger.info(f"Processing document: {file_path.name} ({file_extension})")
        
        if file_extension == '.pdf':
            return await self._process_pdf(file_path)
        elif file_extension == '.docx':
            return await self._process_docx(file_path)
        elif file_extension in ['.txt']:
            return await self._process_text(file_path)
        elif file_extension in ['.png', '.jpg', '.jpeg']:
            return await self._process_image(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_extension}")
    
    async def _process_pdf(self, file_path: Path) -> ExtractedData:
        """Process PDF files - extract text, tables, and images"""
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(self.executor, self._process_pdf_sync, file_path)
    
    def _process_pdf_sync(self, file_path: Path) -> ExtractedData:
        """Synchronous PDF processing"""
        chunks = []
        
        try:
            # Extract text using PyPDF2
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                total_pages = len(pdf_reader.pages)
                
                for page_num, page in enumerate(pdf_reader.pages, 1):
                    text = page.extract_text()
                    if text.strip():
                        chunk_id = str(uuid.uuid4())
                        chunks.append(DocumentChunk(
                            id=chunk_id,
                            text=text,
                            page_number=page_num,
                            content_type=ContentType.TEXT,
                            metadata={"extraction_method": "PyPDF2"}
                        ))
            
            # Try to extract tables using tabula
            try:
                tables = tabula.read_pdf(str(file_path), pages='all', multiple_tables=True)
                for i, table in enumerate(tables):
                    if not table.empty:
                        table_text = table.to_string(index=False)
                        chunk_id = str(uuid.uuid4())
                        chunks.append(DocumentChunk(
                            id=chunk_id,
                            text=f"Table {i+1}:\n{table_text}",
                            page_number=1,  # tabula doesn't provide page info easily
                            content_type=ContentType.TABLE,
                            metadata={"extraction_method": "tabula", "table_index": i}
                        ))
            except Exception as e:
                logger.warning(f"Could not extract tables from PDF: {e}")
            
            return ExtractedData(
                filename=file_path.name,
                chunks=chunks,
                total_pages=total_pages,
                metadata={"processing_method": "pdf_multi"}
            )
            
        except Exception as e:
            logger.error(f"Error processing PDF {file_path}: {e}")
            raise
    
    async def _process_docx(self, file_path: Path) -> ExtractedData:
        """Process DOCX files"""
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(self.executor, self._process_docx_sync, file_path)
    
    def _process_docx_sync(self, file_path: Path) -> ExtractedData:
        """Synchronous DOCX processing"""
        chunks = []
        
        try:
            doc = DocxDocument(file_path)
            
            # Extract paragraphs
            for para_num, para in enumerate(doc.paragraphs):
                if para.text.strip():
                    chunk_id = str(uuid.uuid4())
                    chunks.append(DocumentChunk(
                        id=chunk_id,
                        text=para.text,
                        page_number=1,  # DOCX doesn't have clear page boundaries
                        content_type=ContentType.TEXT,
                        metadata={"paragraph_number": para_num}
                    ))
            
            # Extract tables
            for table_num, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                
                if table_data:
                    table_text = "\n".join(["\t".join(row) for row in table_data])
                    chunk_id = str(uuid.uuid4())
                    chunks.append(DocumentChunk(
                        id=chunk_id,
                        text=f"Table {table_num+1}:\n{table_text}",
                        page_number=1,
                        content_type=ContentType.TABLE,
                        metadata={"table_number": table_num}
                    ))
            
            return ExtractedData(
                filename=file_path.name,
                chunks=chunks,
                total_pages=1,
                metadata={"processing_method": "docx"}
            )
            
        except Exception as e:
            logger.error(f"Error processing DOCX {file_path}: {e}")
            raise
    
    async def _process_text(self, file_path: Path) -> ExtractedData:
        """Process plain text files"""
        chunks = []
        
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
                
                # Split into chunks (simple sentence-based splitting)
                sentences = text.split('.')
                chunk_text = ""
                chunk_num = 1
                
                for sentence in sentences:
                    if len(chunk_text + sentence) > 1000:  # Max chunk size
                        if chunk_text.strip():
                            chunk_id = str(uuid.uuid4())
                            chunks.append(DocumentChunk(
                                id=chunk_id,
                                text=chunk_text.strip(),
                                page_number=chunk_num,
                                content_type=ContentType.TEXT,
                                metadata={"chunk_number": chunk_num}
                            ))
                            chunk_num += 1
                        chunk_text = sentence
                    else:
                        chunk_text += sentence + "."
                
                # Add final chunk
                if chunk_text.strip():
                    chunk_id = str(uuid.uuid4())
                    chunks.append(DocumentChunk(
                        id=chunk_id,
                        text=chunk_text.strip(),
                        page_number=chunk_num,
                        content_type=ContentType.TEXT,
                        metadata={"chunk_number": chunk_num}
                    ))
            
            return ExtractedData(
                filename=file_path.name,
                chunks=chunks,
                total_pages=chunk_num,
                metadata={"processing_method": "text"}
            )
            
        except Exception as e:
            logger.error(f"Error processing text file {file_path}: {e}")
            raise
    
    async def _process_image(self, file_path: Path) -> ExtractedData:
        """Process image files using OCR"""
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(self.executor, self._process_image_sync, file_path)
    
    def _process_image_sync(self, file_path: Path) -> ExtractedData:
        """Synchronous image processing"""
        chunks = []
        
        try:
            # Use OCR to extract text from image
            image = Image.open(file_path)
            text = pytesseract.image_to_string(image)
            
            if text.strip():
                chunk_id = str(uuid.uuid4())
                chunks.append(DocumentChunk(
                    id=chunk_id,
                    text=text,
                    page_number=1,
                    content_type=ContentType.IMAGE,
                    metadata={
                        "extraction_method": "OCR",
                        "image_size": image.size
                    }
                ))
            
            return ExtractedData(
                filename=file_path.name,
                chunks=chunks,
                total_pages=1,
                metadata={"processing_method": "ocr"}
            )
            
        except Exception as e:
            logger.error(f"Error processing image {file_path}: {e}")
            raise